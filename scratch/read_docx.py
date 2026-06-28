import docx
import sys

# Configure stdout to use UTF-8
sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document("SDLC_Model_Industrial_Safety_Chatbot.docx")
print(f"Total paragraphs: {len(doc.paragraphs)}")

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f"[{i}]: {text}")
