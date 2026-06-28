
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(0.8)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── helpers ────────────────────────────────────────────────────────────────────

def set_cell_border(cell, **kwargs):
    """Set individual cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            for k, v in kwargs[edge].items():
                tag.set(qn(f'w:{k}'), v)
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))   # 1 cm ≈ 567 twips
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def merge_cols(row, start, end):
    """Merge cells start..end (0-indexed) in a row."""
    row.cells[start].merge(row.cells[end])

def add_run(para, text, bold=False, size=10, italic=False, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def cell_para(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
              italic=False, color=None, indent_left=0):
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    if indent_left:
        para.paragraph_format.left_indent = Pt(indent_left)
    add_run(para, text, bold=bold, size=size, italic=italic, color=color)
    return para

BORDER_THIN = {'val': 'single', 'sz': '4', 'space': '0', 'color': '000000'}
BORDER_NONE = {'val': 'none',   'sz': '0', 'space': '0', 'color': 'auto'}

def apply_borders(cell, top=True, bottom=True, left=True, right=True):
    kwargs = {}
    kwargs['top']    = BORDER_THIN if top    else BORDER_NONE
    kwargs['bottom'] = BORDER_THIN if bottom else BORDER_NONE
    kwargs['left']   = BORDER_THIN if left   else BORDER_NONE
    kwargs['right']  = BORDER_THIN if right  else BORDER_NONE
    set_cell_border(cell, **kwargs)

# ── Title ──────────────────────────────────────────────────────────────────────

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('CONTENT')
title_run.bold = True
title_run.font.size = Pt(14)
title_run.font.name = 'Times New Roman'
title_para.paragraph_format.space_after = Pt(6)

# ── Table definition ───────────────────────────────────────────────────────────
# Columns: S.no | Heading | (sub blank) | Page No.
# col widths in cm: 1.8 | 7.5 | 4.5 | 2.0

COL_W = [Cm(1.8), Cm(7.5), Cm(4.5), Cm(2.0)]   # 4 cols total

tbl = doc.add_table(rows=0, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

# Fix column widths
for col_idx, width in enumerate(COL_W):
    for cell in tbl.columns[col_idx].cells:
        cell.width = width

def add_row(sno='', heading='', sub='', page='',
            sno_bold=False, heading_bold=False,
            bg=None, row_height=0.55,
            heading_indent=0, sub_indent=0,
            merge_heading_sub=False):
    row = tbl.add_row()
    set_row_height(row, row_height)
    cells = row.cells

    # col widths
    cells[0].width = COL_W[0]
    cells[1].width = COL_W[1]
    cells[2].width = COL_W[2]
    cells[3].width = COL_W[3]

    if bg:
        for c in cells:
            shade_cell(c, bg)

    if merge_heading_sub:
        cells[1].merge(cells[2])

    # S.no cell
    cell_para(cells[0], sno, bold=sno_bold, size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # Heading cell
    cell_para(cells[1], heading, bold=heading_bold, size=10,
              indent_left=heading_indent)

    # Sub cell (only if not merged)
    if not merge_heading_sub:
        cell_para(cells[2], sub, bold=False, size=10,
                  indent_left=sub_indent)

    # Page No
    cell_para(cells[3], str(page) if page != '' else '',
              size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # borders – keep full grid
    for c in cells:
        apply_borders(c)

    return row

# ── Header row ─────────────────────────────────────────────────────────────────

hdr = tbl.add_row()
set_row_height(hdr, 0.6)
shade_cell(hdr.cells[0], 'D9D9D9')
shade_cell(hdr.cells[1], 'D9D9D9')
shade_cell(hdr.cells[2], 'D9D9D9')
shade_cell(hdr.cells[3], 'D9D9D9')
hdr.cells[0].width = COL_W[0]
hdr.cells[1].width = COL_W[1]
hdr.cells[2].width = COL_W[2]
hdr.cells[3].width = COL_W[3]

cell_para(hdr.cells[0], 'S.no',    bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
cell_para(hdr.cells[1], 'Heading', bold=True, size=10)
cell_para(hdr.cells[2], '',        bold=True, size=10)
cell_para(hdr.cells[3], 'Page No.',bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
for c in hdr.cells:
    apply_borders(c)

# ── Data rows ──────────────────────────────────────────────────────────────────
# Format: (sno, heading, sub_col, page, sno_bold, heading_bold, merge_h_s)
# merge_heading_sub=True  → cols 1+2 merged (used for section / subsection headings)

rows_data = [
    # S.no | Heading col | sub col | Page | sno_bold | h_bold | merge
    ('1.',   'INTRODUCTION',       '',    '',    True, True,  True),
    ('',     '1.1',  'Project Description',      '13-15', False, False, False),
    ('',     '1.2',  'Problem Statement',         '',     False, False, False),
    ('',     '1.3',  'Project Objectives',        '',     False, False, False),
    ('',     '1.4',  'Company Profile',          '15-17', False, False, False),

    ('2.',   'SYSTEM STUDY',       '',    '',    True, True,  True),
    ('',     '2.1',  'Existing System vs Proposed System', '18-19', False, False, False),
    ('',     '2.2',  'Feasibility Study',         '',     False, False, False),
    ('',     '',     '2.2.1  Technical Feasibility',   '20', False, False, False),
    ('',     '',     '2.2.2  Economic Feasibility',    '20', False, False, False),
    ('',     '',     '2.2.3  Operational Feasibility', '21', False, False, False),
    ('',     '',     '2.2.4  Schedule Feasibility',    '21', False, False, False),
    ('',     '',     '2.2.5  SDLC Model',              '21-22', False, False, False),
    ('',     '',     '2.2.6  Complexity Factors',      '22-23', False, False, False),
    ('',     '',     '2.2.7  COCOMO Estimation',       '23-24', False, False, False),
    ('',     '2.3',  'Tools and Technologies Used',  '',  False, False, False),
    ('',     '',     '2.3.1  React.js',               '24', False, False, False),
    ('',     '',     '2.3.2  Node.js / Express',      '24-25', False, False, False),
    ('',     '',     '2.3.3  MongoDB',                '25', False, False, False),
    ('',     '',     '2.3.4  OpenCV',                 '25-26', False, False, False),
    ('',     '',     '2.3.5  YOLOv8',                 '26', False, False, False),
    ('',     '',     '2.3.6  NLP / Rasa / Dialogflow','26', False, False, False),
    ('',     '',     '2.3.7  Visual Studio Code (VS Code)', '26-27', False, False, False),
    ('',     '',     '2.3.8  Git and GitHub',          '27', False, False, False),
    ('',     '2.4',  'Hardware and Software Requirements', '', False, False, False),
    ('',     '',     '2.4.1  Hardware Requirements',   '28', False, False, False),
    ('',     '',     '2.4.2  Software Requirements',   '28', False, False, False),

    ('3.',   'SOFTWARE REQUIREMENTS SPECIFICATION', '', '', True, True, True),
    ('',     '3.1',  'User Requirements',             '',  False, False, False),
    ('',     '',     '3.1.1  Administrator',           '28', False, False, False),
    ('',     '',     '3.1.2  Safety Officer',          '29-30', False, False, False),
    ('',     '',     '3.1.3  Employee',                '30', False, False, False),
    ('',     '3.2',  'Functional Requirements',        '',  False, False, False),
    ('',     '',     '3.2.1  User Authentication and Authorization', '30', False, False, False),
    ('',     '',     '3.2.2  Employee Management',     '30-31', False, False, False),
    ('',     '',     '3.2.3  CCTV Video Monitoring',   '31', False, False, False),
    ('',     '',     '3.2.4  PPE Detection Module',    '31', False, False, False),
    ('',     '',     '3.2.5  Employee Identification Module', '31', False, False, False),
    ('',     '',     '3.2.6  Safety Violation Detection', '31-32', False, False, False),
    ('',     '',     '3.2.7  Alert and Notification Management', '32', False, False, False),
    ('',     '',     '3.2.8  Report Generation',       '32', False, False, False),
    ('',     '',     '3.2.9  Dashboard Management',    '32', False, False, False),
    ('',     '',     '3.2.10 Database Management',     '32-33', False, False, False),
    ('',     '3.3',  'Non-Functional Requirements',    '',  False, False, False),
    ('',     '',     '3.3.1  Performance Requirements', '33', False, False, False),
    ('',     '',     '3.3.2  Reliability Requirements', '33', False, False, False),
    ('',     '',     '3.3.3  Security Requirements',   '34', False, False, False),
    ('',     '',     '3.3.4  Usability Requirements',  '34', False, False, False),
    ('',     '',     '3.3.5  Scalability Requirements', '34', False, False, False),
    ('',     '',     '3.3.6  Availability Requirements','34', False, False, False),
    ('',     '',     '3.3.7  Maintainability Requirements', '34-35', False, False, False),
    ('',     '',     '3.3.8  Unadjusted Function Point (UFP) Analysis', '35', False, False, False),

    ('4.',   'SYSTEM DESIGN',      '',    '',    True, True,  True),
    ('',     '4.1',  'System Perspective',            '36-38', False, False, False),
    ('',     '4.2',  'Context Diagram (DFD)',         '',  False, False, False),
    ('',     '',     '4.2.1  DFD Level 0',            '39', False, False, False),
    ('',     '',     '4.2.2  DFD Level 1',            '39-41', False, False, False),
    ('',     '',     '4.2.3  DFD Level 2',            '41-42', False, False, False),
    ('',     '4.3',  'Use Case Diagram',              '',  False, False, False),
    ('',     '',     '4.3.1  Actors of the System',   '43', False, False, False),
    ('',     '',     '4.3.2  Major Use Cases',        '43-44', False, False, False),
    ('',     '',     '4.3.3  Use Case Diagram',       '44', False, False, False),
    ('',     '',     '4.3.4  Use Case Description',   '44-45', False, False, False),
    ('',     '4.4',  'Sequence Diagrams',             '',  False, False, False),
    ('',     '',     '4.4.1  Sequence Diagram for User Login', '47', False, False, False),
    ('',     '',     '4.4.2  Sequence Diagram for Employee Identification', '47-48', False, False, False),
    ('',     '',     '4.4.3  Sequence Diagram for PPE Detection', '48-49', False, False, False),
    ('',     '',     '4.4.4  Sequence Diagram for Safety Violation Alert', '49-50', False, False, False),
    ('',     '4.5',  'Activity Diagram',              '50-51', False, False, False),
    ('',     '4.6',  'Database Design',               '',  False, False, False),
    ('',     '',     '4.6.1  ER Diagram',             '52', False, False, False),
    ('',     '',     '4.6.2  Database Schema',        '52-54', False, False, False),

    ('5.',   'IMPLEMENTATION',     '',    '',    True, True,  True),
    ('',     '5.1',  'Frontend Components',           '55-57', False, False, False),
    ('',     '5.2',  'Backend Architecture',          '57-59', False, False, False),
    ('',     '5.3',  'Database Implementation',       '59-61', False, False, False),

    ('6.',   'TESTING',            '',    '',    True, True,  True),
    ('',     '6.1',  'Test Plan',                     '62', False, False, False),
    ('',     '6.2',  'Test Cases',                    '62-65', False, False, False),
    ('',     '6.3',  'Testing Results',               '65-66', False, False, False),

    ('7.',   'CONCLUSION',         '',    '67',  True, True,  True),
    ('8.',   'FUTURE ENHANCEMENTS','',    '68',  True, True,  True),
    ('',     'BIBLIOGRAPHY',       '',    '69',  False, True, True),
    ('',     'APPENDIX A: Installation Guide', '', '70-72', False, True, True),
    ('',     'APPENDIX B: User Manual',        '', '73-75', False, True, True),
]

for r in rows_data:
    sno, heading, sub, page, sno_bold, h_bold, merge = r
    row = tbl.add_row()
    set_row_height(row, 0.55)
    cells = row.cells

    cells[0].width = COL_W[0]
    cells[1].width = COL_W[1]
    cells[2].width = COL_W[2]
    cells[3].width = COL_W[3]

    is_section = merge and sno != ''
    if is_section:
        shade_cell(cells[0], 'F2F2F2')
        shade_cell(cells[1], 'F2F2F2')
        shade_cell(cells[2], 'F2F2F2')
        shade_cell(cells[3], 'F2F2F2')

    if merge:
        cells[1].merge(cells[2])

    cell_para(cells[0], sno, bold=sno_bold, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_para(cells[1], heading, bold=h_bold, size=10)

    if not merge:
        cell_para(cells[2], sub, size=10)

    cell_para(cells[3], str(page), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    for c in cells:
        apply_borders(c)

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = r'd:\ai chatbot\TABLE_OF_CONTENTS.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
